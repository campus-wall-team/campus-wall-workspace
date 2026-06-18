# -*- coding: utf-8 -*-
"""
校园墙项目商业计划书生成脚本
使用 python-docx 生成 Word 文档
结合创业项目说明文档，突出AI学长功能，4人小团队，精简资金
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def set_run_font(run, font_name='宋体', font_size=12, bold=False, color=None):
    """设置字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    styles = {
        1: {'font': '黑体', 'size': 22, 'color': (0, 51, 102), 'space_before': 24, 'space_after': 12},
        2: {'font': '黑体', 'size': 16, 'color': (0, 51, 102), 'space_before': 18, 'space_after': 10},
        3: {'font': '黑体', 'size': 14, 'color': (0, 51, 102), 'space_before': 14, 'space_after': 8},
    }
    style = styles.get(level, styles[3])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(style['space_before'])
    p.paragraph_format.space_after = Pt(style['space_after'])
    run = p.add_run(text)
    set_run_font(run, font_name=style['font'], font_size=style['size'],
                 bold=True, color=style['color'])
    return p


def add_paragraph_custom(doc, text, indent=True, bold=False, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加自定义段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, font_name='宋体', font_size=font_size, bold=bold)
    return p


def add_bullet_list(doc, items, indent_level=0):
    """添加项目符号列表"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.5)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, font_name='宋体', font_size=11)


def create_cover_page(doc):
    """创建封面页"""
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('校园墙')
    set_run_font(run, font_name='黑体', font_size=36, bold=True, color=(0, 51, 102))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('AI 校园知识助手平台')
    set_run_font(run, font_name='黑体', font_size=28, bold=True, color=(0, 102, 153))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('商业计划书')
    set_run_font(run, font_name='黑体', font_size=32, bold=True, color=(0, 51, 102))

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—— 让校园信息不再沉默，让学长经验代代相传 ——')
    set_run_font(run, font_name='楷体', font_size=14, color=(102, 102, 102))

    doc.add_paragraph()
    doc.add_paragraph()

    info_items = [
        ('项目载体', '嘉应学院校园墙'),
        ('项目类型', '大学生创业 / AI+校园服务'),
        ('核心产品', 'AI学长 · 智能校园知识助手'),
        ('团队规模', '4人创始团队'),
        ('启动资金', '15-20万元'),
        ('文档版本', 'V1.0'),
        ('编制日期', '2026年5月'),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run(f'{label}：')
        set_run_font(run1, font_name='黑体', font_size=12, bold=True)
        run2 = p.add_run(value)
        set_run_font(run2, font_name='宋体', font_size=12)

    doc.add_page_break()


def create_toc_page(doc):
    """创建目录页"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('目  录')
    set_run_font(run, font_name='黑体', font_size=22, bold=True, color=(0, 51, 102))

    toc_items = [
        ('一、项目概述', '3'),
        ('二、市场痛点与机会', '4'),
        ('三、核心产品：AI学长', '6'),
        ('四、技术架构与壁垒', '10'),
        ('五、商业模式', '13'),
        ('六、运营计划', '15'),
        ('七、团队介绍', '17'),
        ('八、财务规划', '18'),
        ('九、发展规划', '20'),
        ('十、所需支持', '21'),
    ]

    for title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run1 = p.add_run(title)
        set_run_font(run1, font_name='宋体', font_size=12)
        tab_run = p.add_run('\t' * 8)
        run2 = p.add_run(page)
        set_run_font(run2, font_name='宋体', font_size=12)

    doc.add_page_break()


def create_project_overview(doc):
    """一、项目概述"""
    add_heading_custom(doc, '一、项目概述', level=1)

    add_heading_custom(doc, '1.1 一句话介绍', level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('我们以"嘉应学院校园墙"为载体，打造一个面向高校学生的智能校园社交与服务平台。')
    set_run_font(run, font_name='楷体', font_size=14, bold=True, color=(0, 102, 153))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run('核心是把分散在微信群、校园墙、学生手册中的海量校园信息，')
    set_run_font(run2, font_name='楷体', font_size=13, color=(0, 102, 153))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(12)
    run3 = p3.add_run('沉淀为一个可对话的"AI学长"——学生用一句话提问，就能得到准确、可溯源的校园答案。')
    set_run_font(run3, font_name='楷体', font_size=13, bold=True, color=(0, 102, 153))

    add_heading_custom(doc, '1.2 项目愿景', level=2)
    add_paragraph_custom(doc,
        '让校园墙从"看一眼就划走的社交流"，升级为"能解决问题的服务入口"。'
        '让每一位大学生都能在一个安全、智能、有温度的社区中，自由地分享生活、获取信息、结识朋友。',
        indent=True)

    add_heading_custom(doc, '1.3 核心差异化定位', level=2)
    add_paragraph_custom(doc,
        '我们做的不是又一个"表白墙"，而是"社交 + 智能服务"双驱动的平台。'
        '把校园里真正有用的信息（办事流程、学长经验、高频问答）沉淀成一个可对话、可溯源的 AI 知识库——"AI学长"，'
        '让校园社交与智能服务融为一体。',
        indent=True)

    add_heading_custom(doc, '1.4 项目现状', level=2)
    status = [
        '已完成可运行的平台原型：社区发帖、即时通讯、AI问答、数据管道、监控告警全栈就绪',
        '已沉淀嘉应学院体测、教务等真实问答数据，形成冷启动数据壁垒',
        '技术架构完整：前后端分离 + 微服务 + GraphRAG知识图谱 + 多端支持',
        '现计划组建4人创始团队、注册公司，先深耕嘉应学院，再逐步复制到更多高校',
    ]
    add_bullet_list(doc, status)


def create_market_analysis(doc):
    """二、市场痛点与机会"""
    add_heading_custom(doc, '二、市场痛点与机会', level=1)

    add_heading_custom(doc, '2.1 现有校园墙的三大硬伤', level=2)

    pain_points = [
        ('内容"发完即沉"', '信息靠人工转发，发布后便淹没在时间线里，无法检索、容易过时。新生问体测免测、证件补办、教务与宿舍流程，答案散落在微信群历史、公告和学长经验里，搜索困难、重复提问。'),
        ('只有"社交"没有"服务"', '现有校园墙功能局限于匿名表白、吐槽、失物招领、二手交易，解决不了学生真正高频的办事与答疑需求。校园社交产品普遍缺少"智能答疑"这一学生刚需能力。'),
        ('经验无法沉淀', '学长经验靠口口相传，学生一毕业就流失，届届重复提问。宝贵经验没有系统化沉淀，每一届新生都要从头摸索。'),
    ]

    for title, desc in pain_points:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.74)
        run1 = p.add_run(f'{title}：')
        set_run_font(run1, font_name='黑体', font_size=11, bold=True, color=(180, 0, 0))
        run2 = p.add_run(desc)
        set_run_font(run2, font_name='宋体', font_size=11)

    add_heading_custom(doc, '2.2 数据合规隐患', level=2)
    add_paragraph_custom(doc,
        '直接拿聊天数据做 AI 存在隐私与合规隐患，市面缺乏合规可用的解决方案。'
        '如何在利用数据提升 AI 能力的同时保护学生隐私，是一个尚未被解决的关键问题。',
        indent=True)

    add_heading_custom(doc, '2.3 市场机会', level=2)
    opportunities = [
        '中国高等教育在校生总数超过 4700 万人，每年新增入学人数约 1000 万，市场基数巨大',
        '高校学生智能手机普及率超过 99%，微信小程序日均使用时长持续增长',
        '校园生活服务市场规模预计 2026 年突破 8000 亿元',
        'AI 技术普及降低了智能服务的门槛，为校园场景的创新提供了技术基础',
        '单校深耕模式验证成本低，成功后可快速复制到其他高校',
    ]
    add_bullet_list(doc, opportunities)

    add_heading_custom(doc, '2.4 目标用户', level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['用户类型', '典型场景', '核心需求']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_name='黑体', font_size=11, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    user_data = [
        ('新生', '体测免测怎么申请？宿舍怎么选？校园卡在哪补办？', '快速获取准确的校园办事信息'),
        ('在校生', '奖学金评定标准？综测加分活动有哪些？选课攻略？', '了解学校政策和学长经验'),
        ('活动组织者', '团学活动如何高效触达学生？如何提升参与度？', '精准推送活动信息，提高覆盖面'),
        ('校园商家', '如何精准触达有需求的学生？', '精准营销，降低获客成本'),
    ]

    for user_type, scene, need in user_data:
        row_cells = table.add_row().cells
        row_cells[0].text = user_type
        row_cells[1].text = scene
        row_cells[2].text = need
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=10)

    doc.add_paragraph()


def create_ai_product(doc):
    """三、核心产品：AI学长"""
    add_heading_custom(doc, '三、核心产品：AI学长', level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('AI学长 —— 你的24小时在线校园助手')
    set_run_font(run, font_name='楷体', font_size=16, bold=True, color=(0, 102, 153))

    add_paragraph_custom(doc,
        'AI学长是校园墙平台的核心功能，它将分散在校园各处的海量信息沉淀为可对话、可溯源的 AI 知识库。'
        '学生用一句话提问，就能得到准确、有依据的校园答案。',
        indent=True)

    add_heading_custom(doc, '3.1 两大信息场景', level=2)

    add_paragraph_custom(doc, '（1）长期信息 —— 知识沉淀', indent=True, bold=True)
    long_term = [
        '将学校规章、办事流程、学长经验等基础信息导入小程序知识库',
        '学生有疑问可直接询问"AI学长"，由知识库检索给出准确答案',
        '答案可溯源，每条回答都标注信息来源，确保可靠性',
        '支持图片回复：新生想要了解学校环境，通过拍照学校环境构建大量知识存储，AI学长可以直接将图片发送给学生',
    ]
    add_bullet_list(doc, long_term)

    add_paragraph_custom(doc, '（2）时效信息 —— 活动直达', indent=True, bold=True)
    short_term = [
        '痛点：团学活动层层下发（团学 → 部门 → 班委 → 学生），易出现漏发、错发',
        '方案：团学活动直接发布到 AI 校园墙并触达学生',
        '智能推荐：当学生咨询奖学金时，AI学长主动给出方案——推荐综测加分高的活动，引导学生争取奖学金，提升积极性与主动性',
        '个性化推送：根据学生画像，推送与其相关的活动和机会',
    ]
    add_bullet_list(doc, short_term)

    add_heading_custom(doc, '3.2 AI学长的技术实现', level=2)

    add_paragraph_custom(doc,
        'AI学长不是简单关键词搜索或泛泛大模型回答，而是采用"向量检索 + 知识图谱"双通道架构，'
        '答案结构化、可溯源、更准确。',
        indent=True)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['技术模块', '实现方式', '作用']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_name='黑体', font_size=10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tech_data = [
        ('GraphRAG 知识图谱', 'Neo4j + 向量检索 + 图遍历', '将校园信息构建为结构化知识图谱，实现精准关联问答'),
        ('双 LLM 架构', 'DashScope/qwen-plus + Ollama/bge-m3', 'Chat LLM负责推理回答，Embedding LLM负责向量化，兼顾效果与成本'),
        ('数据管道 ETL', 'Python ETL + 自动脱敏', '从真实校园数据中提取知识，自动清洗、脱敏后构建知识库'),
        ('多轮对话', '上下文记忆 + 会话管理', '支持连续追问，理解上下文，提供更自然的交互体验'),
        ('图片知识库', '图片索引 + 多模态检索', '校园环境照片、办事流程截图等以图搜图、以文搜图'),
    ]

    for module, impl, effect in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = module
        row_cells[1].text = impl
        row_cells[2].text = effect
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=10)

    doc.add_paragraph()

    add_heading_custom(doc, '3.3 AI学长的典型使用场景', level=2)

    scenarios = [
        ('新生入学', '问："宿舍怎么选？" → AI学长：给出选房流程、注意事项，并附带宿舍环境图片'),
        ('办事指南', '问："体测免测怎么申请？" → AI学长：给出申请条件、所需材料、办理地点和截止时间'),
        ('活动参与', '问："怎么拿奖学金？" → AI学长：解释评定标准，并推荐当前可参加的综测加分活动'),
        ('学长经验', '问："选课有什么建议？" → AI学长：汇总历届学长经验，给出分学期选课策略'),
        ('校园导航', '问："图书馆几点开门？" → AI学长：给出开放时间，并推送图书馆近期活动'),
        ('紧急求助', '问："校园卡丢了怎么办？" → AI学长：给出补办流程、地点、所需材料和费用'),
    ]

    for title, desc in scenarios:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.74)
        run1 = p.add_run(f'{title}：')
        set_run_font(run1, font_name='黑体', font_size=11, bold=True)
        run2 = p.add_run(desc)
        set_run_font(run2, font_name='宋体', font_size=11)

    add_heading_custom(doc, '3.4 社交功能矩阵', level=2)
    add_paragraph_custom(doc,
        '除AI学长外，平台还具备完整的校园社交功能，与AI服务形成互补：',
        indent=True)

    social = [
        '圈子（帖子浏览与发布）：匿名发帖、分类浏览（失物招领、二手交易、表白墙、求助问答）、点赞评论互动、热帖排行',
        '私信系统：用户间实时私信聊天，基于 WebSocket 的即时通讯，消息通知与未读提醒',
        '用户系统：微信一键登录（OAuth2）、用户画像与偏好设置、浏览历史与收藏管理、关注/粉丝关系',
        '内容审核：AI自动审核 + 人工复审双机制，敏感词过滤，违规内容自动标记，管理后台审核队列',
    ]
    add_bullet_list(doc, social)


def create_tech_barrier(doc):
    """四、技术架构与壁垒"""
    add_heading_custom(doc, '四、技术架构与壁垒', level=1)

    add_heading_custom(doc, '4.1 系统架构概览', level=2)
    add_paragraph_custom(doc,
        '校园墙采用前后端分离 + 微服务架构，由 6 个独立服务模块组成，确保系统的高可用、高扩展和高安全性。',
        indent=True)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['层次', '服务', '技术栈', '职责']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_name='黑体', font_size=10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    arch_data = [
        ('客户端', 'H5/小程序/管理后台', 'Vue 3 + uni-app + Element Plus', '用户交互，一套代码多端运行'),
        ('网关层', 'Nginx', 'Nginx 1.27', '反向代理、负载均衡、SSL终端'),
        ('核心业务', 'campus_wall', 'Java 17 + Spring Boot 3', '用户认证、帖子、评论、AI代理、私信、审核'),
        ('AI引擎', 'campus-wall-ai', 'Python + FastAPI + Neo4j', '知识图谱构建、AI问答、向量检索'),
        ('数据处理', 'campus-wall-data-pipeline', 'Python ETL', '数据提取、清洗、脱敏、推送至知识库'),
        ('数据存储', 'MySQL + Redis + MinIO', 'MySQL 8.0 + Redis 7.0', '业务数据、缓存、图片存储'),
        ('知识存储', 'Neo4j', 'Neo4j 5', '知识图谱 + 原生向量索引'),
        ('监控告警', 'Prometheus + Grafana', 'Prometheus 2.54 + Grafana 11.2', '全链路监控、可视化、告警通知'),
    ]

    for layer, service, tech, duty in arch_data:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = service
        row_cells[2].text = tech
        row_cells[3].text = duty
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=9)

    doc.add_paragraph()

    add_heading_custom(doc, '4.2 四大核心壁垒', level=2)

    barriers = [
        ('GraphRAG 知识图谱问答',
         '不是简单关键词搜索或泛泛大模型，而是"向量检索 + 知识图谱"双通道架构。'
         '答案结构化、可溯源、更准确，这是市面上通用大模型无法比拟的校园垂直优势。'),
        ('合规数据管道（核心壁垒）',
         '自研脱敏与假名化管道，把真实校园数据安全转化为知识库。'
         '自动移除手机号、身份证号、银行卡号等 PII 信息，实测隐私信息零残留。'
         '这是别人不敢碰或做不合规的护城河。'),
        ('校园垂直深耕',
         '聚焦单校真实场景，已沉淀嘉应学院体测、教务等真实问答数据，形成冷启动数据壁垒。'
         '每一所高校的知识库都是独特的，先发优势明显。'),
        ('平台化可复制',
         '社区 + 即时通讯 + AI 一体化，向量后端可热切换，工程完整。'
         '验证单校模型后，可快速复制到新学校，边际成本递减。'),
    ]

    for title, desc in barriers:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(0.74)
        run1 = p.add_run(f'{title}\n')
        set_run_font(run1, font_name='黑体', font_size=12, bold=True, color=(0, 102, 153))
        run2 = p.add_run(desc)
        set_run_font(run2, font_name='宋体', font_size=11)

    add_heading_custom(doc, '4.3 安全与合规', level=2)
    security = [
        '认证层：微信 OAuth2 授权 + JWT Token，无密码泄露风险',
        '传输层：全站 HTTPS 加密',
        '数据层：PII 数据自动脱敏、加密存储、定期备份',
        '内容层：敏感词过滤 + AI复审 + 人工审核三重防护',
        '监控层：异常行为检测、安全告警、审计日志',
    ]
    add_bullet_list(doc, security)


def create_business_model(doc):
    """五、商业模式"""
    add_heading_custom(doc, '五、商业模式', level=1)

    add_paragraph_custom(doc,
        '校园墙采用"免费基础服务 + 精准信息推送 + 校园服务订阅"的多元商业模式，'
        '在保障用户体验的前提下实现可持续盈利。',
        indent=True)

    add_heading_custom(doc, '5.1 AI学长精准信息推送', level=2)
    add_paragraph_custom(doc,
        '在"AI学长"中接入精准信息推送，当学生提出相关需求时，AI学长智能推荐匹配服务，平台从撮合中收取少量手续费。',
        indent=True)

    push_examples = [
        '学生找驾校 → AI学长推送驾校列表与联系方式 → 平台收取推荐手续费',
        '学生找兼职 → AI学长推送匹配岗位信息 → 平台收取撮合佣金',
        '学生咨询奖学金 → AI学长推荐综测加分高的活动 → 活动方支付推广费',
        '学生需要校园服务 → 接入校园卡等校园服务 → 服务方支付接入费',
    ]
    add_bullet_list(doc, push_examples)

    add_heading_custom(doc, '5.2 校园服务订阅', level=2)
    add_paragraph_custom(doc,
        '面向院系、社团、校园商家的服务订阅模式：',
        indent=True)

    subscription = [
        '院系服务订阅：为各院系提供定制化的AI问答服务和活动发布平台',
        '社团推广服务：社团活动精准推送给感兴趣的学生，按效果付费',
        '商家入驻服务：校园周边商家入驻平台，获取精准流量',
    ]
    add_bullet_list(doc, subscription)

    add_heading_custom(doc, '5.3 增值服务（To C）', level=2)
    c_services = [
        'AI学长高级版：更多轮次对话、更快速度、更详细的回答',
        '帖子置顶推广：用户付费置顶自己的帖子，获取更高曝光',
        '会员体系：享受去广告、专属标识、优先客服等权益',
    ]
    add_bullet_list(doc, c_services)

    add_heading_custom(doc, '5.4 收入预测', level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['收入来源', '第1年', '第2年', '第3年']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_name='黑体', font_size=10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    revenue_data = [
        ('精准信息推送', '2万', '10万', '30万'),
        ('校园服务订阅', '1万', '8万', '25万'),
        ('增值服务', '0.5万', '5万', '20万'),
        ('广告收入', '0.5万', '5万', '15万'),
        ('收入合计', '4万', '28万', '90万'),
    ]

    for source, y1, y2, y3 in revenue_data:
        row_cells = table.add_row().cells
        row_cells[0].text = source
        row_cells[1].text = y1
        row_cells[2].text = y2
        row_cells[3].text = y3
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def create_operation_plan(doc):
    """六、运营计划"""
    add_heading_custom(doc, '六、运营计划', level=1)

    add_heading_custom(doc, '6.1 三阶段推进策略', level=2)

    phases = [
        ('第一阶段：单校验证期（0-6个月）',
         '目标：深耕嘉应学院，验证产品模型，积累首批 3000+ 用户',
         [
             '与学校团委、学生会建立合作，获取官方背书',
             '在新生入学季集中推广，快速获取种子用户',
             '高频内容运营：每日精选优质内容、策划校园热点话题',
             '完善AI学长知识库：持续录入学校规章、办事流程、学长经验',
             '建立校园大使网络：每个学院招募2-3名校园大使',
             '里程碑：日活500+、AI问答可用率>90%、内容审核流程跑通',
         ]),
        ('第二阶段：校内扩张期（6-12个月）',
         '目标：覆盖嘉应学院全校，用户规模达到 8000+',
         [
             '拓展服务场景：接入更多校园服务（校园卡、图书馆、教务等）',
             '启动商业化试点：精准信息推送、商家入驻',
             '优化产品体验：根据用户反馈持续迭代AI学长功能',
             '建立数据壁垒：持续沉淀校园知识，形成不可替代性',
             '里程碑：日活2000+、实现初步商业化收入、盈亏平衡',
         ]),
        ('第三阶段：多校复制期（12-24个月）',
         '目标：复制到3-5所高校，用户规模达到3万+',
         [
             '标准化复制流程：知识库迁移、校园大使培训、推广SOP',
             '优先复制同城高校：利用地缘优势降低推广成本',
             '探索SaaS模式：为其他学校提供校园墙平台部署服务',
             '里程碑：覆盖5所高校、年收入50万+、实现盈利',
         ]),
    ]

    for title, goal, items in phases:
        add_heading_custom(doc, title, level=3)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.74)
        run = p.add_run(f'目标：{goal}')
        set_run_font(run, font_name='黑体', font_size=11, bold=True, color=(0, 102, 153))
        add_bullet_list(doc, items)

    add_heading_custom(doc, '6.2 用户增长策略', level=2)
    growth = [
        '新生迎新：在新生报到期间设置推广点，引导注册并使用AI学长',
        'KOL合作：邀请校园公众号、校园红人进行产品体验推荐',
        '口碑裂变：邀请奖励机制，邀请新用户双方获得积分/特权',
        '活动营销：举办"最美校园瞬间"摄影大赛等线上活动吸引用户',
        '社群运营：建立校园兴趣圈子，增强用户归属感',
    ]
    add_bullet_list(doc, growth)

    add_heading_custom(doc, '6.3 用户留存策略', level=2)
    retention = [
        'AI服务粘性：AI学长的实用价值使用户形成使用依赖',
        '内容运营：每日精选优质内容推送，保持用户打开习惯',
        '消息推送：个性化消息推送，召回沉默用户',
        '互动激励：签到、发帖、评论等行为获得积分和等级提升',
    ]
    add_bullet_list(doc, retention)


def create_team(doc):
    """七、团队介绍"""
    add_heading_custom(doc, '七、团队介绍', level=1)

    add_paragraph_custom(doc,
        '我们是一支4人的大学生创业团队，技术背景扎实，对校园场景有深刻理解。'
        '团队成员分工明确，协作高效，具备从0到1构建产品的完整能力。',
        indent=True)

    add_heading_custom(doc, '7.1 创始团队（4人）', level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['角色', '人数', '核心职责', '要求']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_name='黑体', font_size=10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    team_data = [
        ('技术负责人', '1人', '后端开发、系统架构、AI服务集成', 'Java/Python开发能力，熟悉Spring Boot/FastAPI'),
        ('AI工程师', '1人', 'GraphRAG知识图谱、LLM调优、数据管道', 'Python开发，了解机器学习/大模型应用'),
        ('前端工程师', '1人', 'uni-app小程序开发、管理后台、UI实现', 'Vue 3/TypeScript，熟悉uni-app跨端开发'),
        ('运营负责人', '1人', '内容运营、用户增长、商务合作、市场推广', '沟通能力强，熟悉校园推广渠道'),
    ]

    for role, count, duty, req in team_data:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = count
        row_cells[2].text = duty
        row_cells[3].text = req
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=10)

    doc.add_paragraph()

    add_heading_custom(doc, '7.2 团队优势', level=2)
    advantages = [
        '全栈技术能力：覆盖后端、前端、AI、运维全链路，无需外包',
        '校园场景理解：团队成员本身就是大学生，深刻理解校园痛点',
        '产品已验证：平台原型已运行，技术风险低',
        '轻资产运营：4人小团队，沟通成本低，决策效率高',
        '持续迭代能力：团队具备快速响应市场反馈、持续迭代产品的能力',
    ]
    add_bullet_list(doc, advantages)

    add_heading_custom(doc, '7.3 顾问资源', level=2)
    advisors = [
        '技术顾问：邀请有经验的工程师提供架构指导和技术选型建议',
        '创业导师：依托学校创业孵化基地，获取创业指导和资源对接',
        '教育顾问：高校教育工作者，提供校园市场洞察和政策指导',
    ]
    add_bullet_list(doc, advisors)


def create_financial(doc):
    """八、财务规划"""
    add_heading_custom(doc, '八、财务规划', level=1)

    add_heading_custom(doc, '8.1 启动资金需求：15-20万元', level=2)
    add_paragraph_custom(doc,
        '作为大学生创业项目，我们坚持精益创业理念，控制成本，用最小资金验证商业模式。',
        indent=True)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['用途', '金额（万元）', '占比', '说明']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_name='黑体', font_size=10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fund_data = [
        ('服务器/算力/大模型接口', '6-8', '40%', '云服务器、数据库、CDN、AI算力费用'),
        ('团队运营', '4-6', '30%', '4人基本生活补贴、市场推广物料'),
        ('公司注册/合规法务', '2-3', '15%', '工商注册、法律咨询、合规审查'),
        ('产品迭代与备用金', '3-4', '15%', '产品优化、应急储备、活动经费'),
        ('合计', '15-20', '100%', '种子期启动资金'),
    ]

    for usage, amount, ratio, desc in fund_data:
        row_cells = table.add_row().cells
        row_cells[0].text = usage
        row_cells[1].text = amount
        row_cells[2].text = ratio
        row_cells[3].text = desc
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_heading_custom(doc, '8.2 资金使用计划（第1年）', level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['季度', '主要支出', '金额（万元）', '里程碑']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_name='黑体', font_size=10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    quarter_data = [
        ('Q1', '服务器部署、公司注册、产品开发', '4-5', '平台上线、公司成立'),
        ('Q2', '市场推广、运营活动、产品迭代', '4-5', '日活500+、AI学长可用'),
        ('Q3', '功能优化、商业化试点、团队补贴', '3-4', '日活1500+、首笔收入'),
        ('Q4', '持续运营、知识库扩充、复盘规划', '3-4', '日活3000+、盈亏平衡'),
    ]

    for quarter, expense, amount, milestone in quarter_data:
        row_cells = table.add_row().cells
        row_cells[0].text = quarter
        row_cells[1].text = expense
        row_cells[2].text = amount
        row_cells[3].text = milestone
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_heading_custom(doc, '8.3 财务预测（三年）', level=2)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['指标', '第1年', '第2年', '第3年', '说明']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, font_name='黑体', font_size=10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    forecast_data = [
        ('收入', '4万', '28万', '90万', '精准推送+订阅+增值'),
        ('成本', '18万', '25万', '50万', '服务器+团队+推广'),
        ('净利润', '-14万', '3万', '40万', '第2年实现盈利'),
        ('用户规模', '3000', '15000', '50000', '覆盖高校数'),
        ('覆盖高校', '1所', '3所', '5所', '嘉应学院起步'),
    ]

    for metric, y1, y2, y3, note in forecast_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = y1
        row_cells[2].text = y2
        row_cells[3].text = y3
        row_cells[4].text = note
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_paragraph_custom(doc,
        '说明：以上预测基于保守估计。作为精益创业项目，我们控制初期投入，优先验证商业模式。'
        '第1年以单校深耕为主，第2年开始复制扩张并寻求A轮融资。',
        indent=True)


def create_development_plan(doc):
    """九、发展规划"""
    add_heading_custom(doc, '九、发展规划', level=1)

    add_heading_custom(doc, '9.1 近期目标（6个月）', level=2)
    short_term = [
        '完成公司注册，入驻学校创业孵化基地',
        '平台正式上线运营，日活达到 500+',
        'AI学长知识库覆盖嘉应学院主要办事流程和常见问题',
        '与学校团委、学生会建立正式合作关系',
        '建立校园大使网络，覆盖主要学院',
    ]
    add_bullet_list(doc, short_term)

    add_heading_custom(doc, '9.2 中期目标（12个月）', level=2)
    mid_term = [
        '覆盖嘉应学院全校，用户规模达到 8000+，日活 2000+',
        '启动商业化试点，实现初步收入',
        '启动A轮融资（50-100万元），支撑多校复制',
        '标准化复制流程，完成第2所高校的落地',
    ]
    add_bullet_list(doc, mid_term)

    add_heading_custom(doc, '9.3 长期目标（24个月）', level=2)
    long_term = [
        '覆盖 5 所高校，用户规模达到 5 万+',
        '实现年收入 90 万+，净利润转正',
        '探索SaaS模式，为其他学校提供校园墙平台部署服务',
        '成为粤东地区高校社交与信息服务领先平台',
    ]
    add_bullet_list(doc, long_term)

    add_heading_custom(doc, '9.4 愿景', level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('让每一所高校都有自己的"AI学长"\n')
    set_run_font(run, font_name='楷体', font_size=16, bold=True, color=(0, 102, 153))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('让校园信息不再沉默，让学长经验代代相传')
    set_run_font(run2, font_name='楷体', font_size=14, bold=True, color=(0, 102, 153))


def create_support(doc):
    """十、所需支持"""
    add_heading_custom(doc, '十、所需支持', level=1)

    add_paragraph_custom(doc,
        '作为嘉应学院的大学生创业项目，我们期待在以下方面获得学校的支持，'
        '共同推动校园信息化建设，服务广大师生。',
        indent=True)

    supports = [
        ('启动资金',
         '用于服务器与算力（推理、向量库）、大模型接口费用、团队基本运营、公司注册与合规法务。'
         '我们已制定详细的资金使用计划，确保每一分钱都用在刀刃上。'),
        ('办公场地',
         '希望入驻学校创业孵化基地 / 大学科技园，提供工位与基础办公条件。'
         '降低创业初期成本，同时方便与学校各部门沟通协作。'),
        ('资源对接',
         '希望学校在数据授权、官方渠道试点推广、创业导师指导上给予支持。'
         '学校官方背书将极大提升学生信任度，加速用户增长。'),
        ('政策扶持',
         '申请大学生创业补贴、工商税务绿色通道、创业贷款贴息等政策扶持。'
         '减轻创业初期的合规成本和资金压力。'),
    ]

    for title, desc in supports:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Cm(0.74)
        run1 = p.add_run(f'{title}\n')
        set_run_font(run1, font_name='黑体', font_size=12, bold=True, color=(0, 102, 153))
        run2 = p.add_run(desc)
        set_run_font(run2, font_name='宋体', font_size=11)

    add_heading_custom(doc, '10.1 学校可获得的回报', level=2)
    returns = [
        '校园信息化标杆：打造嘉应学院智慧校园的示范案例',
        '学生服务升级：为学生提供7x24小时智能问答服务，减轻学校咨询压力',
        '数据安全可控：自研合规数据管道，确保校园数据安全可控',
        '就业带动：为在校生提供实习和就业机会，培养技术人才',
        '品牌影响力：学校支持大学生创业，提升学校创新创业教育影响力',
    ]
    add_bullet_list(doc, returns)

    # 结尾
    doc.add_paragraph()
    doc.add_paragraph()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_end = p_end.add_run('—— 感谢阅读，期待合作 ——')
    set_run_font(run_end, font_name='楷体', font_size=14, color=(153, 153, 153))


def main():
    """主函数：生成商业计划书"""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    create_cover_page(doc)
    create_toc_page(doc)
    create_project_overview(doc)
    create_market_analysis(doc)
    create_ai_product(doc)
    create_tech_barrier(doc)
    create_business_model(doc)
    create_operation_plan(doc)
    create_team(doc)
    create_financial(doc)
    create_development_plan(doc)
    create_support(doc)

    output_path = r'E:\Code\project\campus-wall-workspace\business-plan\校园墙商业计划书.docx'
    doc.save(output_path)
    print(f'商业计划书已生成：{output_path}')


if __name__ == '__main__':
    main()
